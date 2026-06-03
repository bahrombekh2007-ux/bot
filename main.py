import asyncio
import copy
import json
import os
import random
import re
import time
from aiohttp import web
from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart, Command
from aiogram.types import (
    Message,
    ReplyKeyboardMarkup,
    KeyboardButton,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    CallbackQuery,
)
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from datetime import datetime, timezone, timedelta

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ─────────────────────────── SOZLAMALAR ──────────────────────────────────────

TOKEN = os.getenv("BOT_TOKEN")
if not TOKEN:
    raise ValueError("BOT_TOKEN muhit o'zgaruvchisi o'rnatilmagan! .env fayliga yoki Render Environment ga qo'shing.")

DATA_FILE = "data/users.json"
os.makedirs("temp", exist_ok=True)
os.makedirs("data", exist_ok=True)

UZBEKISTAN_TZ = timezone(timedelta(hours=5))
SUPPORTED_EXT = (".docx", ".doc", ".txt", ".xlsx", ".pdf")

bot = Bot(token=TOKEN)
dp = Dispatcher()


# ─────────────────────────── YORDAMCHI FUNKSIYALAR ───────────────────────────

def get_uz_time():
    return datetime.now(UZBEKISTAN_TZ)


def get_uz_time_str():
    return get_uz_time().strftime("%d.%m.%Y %H:%M")


def cleanup_old_files(directory="temp", max_age_days=3):
    if not os.path.exists(directory):
        return
    cutoff = time.time() - max_age_days * 86400
    for name in os.listdir(directory):
        path = os.path.join(directory, name)
        try:
            if os.path.isfile(path) and os.path.getmtime(path) < cutoff:
                os.remove(path)
        except Exception:
            pass


# ─────────────────────────── PERSISTENCE ─────────────────────────────────────

def load_users() -> dict:
    """Foydalanuvchi ma'lumotlarini fayldan yuklaydi."""
    try:
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        print(f"⚠️ Ma'lumot yuklashda xatolik: {e}")
    return {}


def save_users(data: dict):
    """Foydalanuvchi ma'lumotlarini faylga saqlaydi."""
    try:
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
    except Exception as e:
        print(f"⚠️ Ma'lumot saqlashda xatolik: {e}")


# Foydalanuvchi ma'lumotlari (runtime)
users: dict = load_users()
user_messages: dict = {}


def get_user(uid: int) -> dict:
    """Foydalanuvchi ma'lumotini qaytaradi yoki yangi yaratadi."""
    key = str(uid)
    if key not in users:
        users[key] = {
            "first_visit": get_uz_time_str(),
            "total_tests": 0,
            "total_questions": 0,
            "total_correct": 0,
            "results": [],
            "uploaded_docs": [],
        }
        save_users(users)
    return users[key]


def save_user(uid: int):
    """Bitta foydalanuvchi ma'lumotini saqlaydi."""
    save_users(users)


# ─────────────────────────── PARSERLAR ───────────────────────────────────────

def _read_txt(file_path: str) -> str:
    """TXT faylni turli kodlashlarda o'qiydi."""
    for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            pass
    return ""


def _parse_hash_format(lines: list) -> list:
    """
    FORMAT A — # va savol bir qatorda:
        # Savol matni
        + To'g'ri javob
        - Noto'g'ri 1

    FORMAT B — # yolg'iz qatorda:
        #
        Savol matni
        +
        To'g'ri javob
    """
    questions = []
    current_q = None
    correct = None
    opts = []
    state = "idle"

    def flush():
        nonlocal current_q, correct, opts, state
        if current_q and correct and len(opts) >= 2:
            all_opts = opts[:4]
            while len(all_opts) < 4:
                all_opts.append(f"Variant {len(all_opts) + 1}")
            if correct not in all_opts:
                all_opts.insert(0, correct)
                all_opts = all_opts[:4]
            questions.append({
                "question": current_q,
                "options": all_opts,
                "answer": correct,
            })
        current_q, correct, opts, state = None, None, [], "idle"

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("#") or line.startswith("?"):
            flush()
            q = line[1:].strip()
            q = re.sub(r"\s*\?$", "", q).strip()
            if q:
                current_q = q
                state = "idle"
            else:
                state = "need_q"

        elif state == "need_q":
            current_q = re.sub(r"\s*\?$", "", line).strip()
            state = "idle"

        elif line.startswith("+"):
            ans = line[1:].strip()
            if not ans:
                state = "need_correct"
            elif current_q is not None:
                correct = ans
                if ans not in opts:
                    opts.append(ans)
                state = "idle"

        elif state == "need_correct":
            correct = line
            if line not in opts:
                opts.append(line)
            state = "idle"

        elif line.startswith("-"):
            ans = line[1:].strip()
            if not ans:
                state = "need_wrong"
            elif current_q is not None and ans not in opts:
                opts.append(ans)

        elif state == "need_wrong":
            if current_q is not None and line not in opts:
                opts.append(line)
            state = "idle"

    flush()
    return questions


def _parse_numbered_abcd(lines: list) -> list:
    """
    1. Savol matni
    A) To'g'ri javob
    B) Noto'g'ri 1
    ...
    Javob: A
    """
    questions = []
    current_q = None
    opts_dict = {}
    correct_letter = None

    def flush():
        nonlocal current_q, opts_dict, correct_letter
        if current_q and opts_dict and correct_letter:
            ul = correct_letter.upper()
            if ul in opts_dict:
                correct_ans = opts_dict[ul]
                options = list(opts_dict.values())[:4]
                while len(options) < 4:
                    options.append(f"Variant {len(options) + 1}")
                questions.append({
                    "question": current_q,
                    "options": options,
                    "answer": correct_ans,
                })
        current_q, opts_dict, correct_letter = None, {}, None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        m = re.match(r"^(\d+)[.)]\s+(.+)$", line)
        if m:
            flush()
            current_q = m.group(2).strip()
            continue

        m = re.match(r"^([A-Da-d])[.)]\s+(.+)$", line)
        if m:
            opts_dict[m.group(1).upper()] = m.group(2).strip()
            continue

        m = re.match(
            r"^(?:Javob|To'g'ri\s*javob|Answer|Ans|Togri\s*javob|Javobi)[:\s]*([A-Da-d])",
            line, re.IGNORECASE,
        )
        if m:
            correct_letter = m.group(1).upper()

    flush()
    return questions


def _parse_pipe(lines: list) -> list:
    """Savol|Javob A|Javob B|Javob C|Javob D"""
    questions = []
    for line in lines:
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 5 and parts[0]:
            questions.append({
                "question": parts[0],
                "options": parts[1:5],
                "answer": parts[1],
            })
    return questions


def _parse_qa_format(lines: list) -> list:
    """
    Savol: Savol matni?
    Javob: To'g'ri javob
    """
    questions = []
    current_q = None
    current_ans = None

    def flush():
        nonlocal current_q, current_ans
        if current_q and current_ans:
            questions.append({
                "question": current_q,
                "options": [current_ans, "Variant B", "Variant C", "Variant D"],
                "answer": current_ans,
            })
        current_q, current_ans = None, None

    for line in lines:
        line = line.strip()
        if not line:
            flush()
            continue

        if re.match(r"^(savol|question)\s*:", line, re.IGNORECASE):
            flush()
            current_q = re.split(r":\s*", line, maxsplit=1)[-1].strip()
        elif re.match(r"^(javob|answer)\s*:", line, re.IGNORECASE):
            current_ans = re.split(r":\s*", line, maxsplit=1)[-1].strip()

    flush()
    return questions


def parse_txt(file_path: str) -> list:
    """TXT faylni avtomatik format aniqlab parse qiladi."""
    content = _read_txt(file_path)
    if not content:
        return []

    lines = content.splitlines()
    non_empty = [l.strip() for l in lines if l.strip()]
    if not non_empty:
        return []

    has_hash = any(l.startswith("#") for l in non_empty)
    has_question_mark = any(l.startswith("?") for l in non_empty)
    has_plus = any(l.startswith("+") for l in non_empty)
    has_pipe = any("|" in l and l.count("|") >= 4 for l in non_empty)
    has_numbered = any(re.match(r"^\d+[.)]\s+", l) for l in non_empty)
    has_abcd = any(re.match(r"^[A-Da-d][.)]\s+", l) for l in non_empty)
    has_javob_key = any(re.match(r"^(javob|answer)\s*:", l, re.IGNORECASE) for l in non_empty)
    has_savol_key = any(re.match(r"^(savol|question)\s*:", l, re.IGNORECASE) for l in non_empty)

    if (has_hash or has_question_mark) and has_plus:
        result = _parse_hash_format(lines)
        if result:
            return result

    if has_numbered and has_abcd:
        result = _parse_numbered_abcd(lines)
        if result:
            return result

    if has_pipe:
        result = _parse_pipe(non_empty)
        if result:
            return result

    if has_savol_key or has_javob_key:
        result = _parse_qa_format(lines)
        if result:
            return result

    for parser in [_parse_hash_format, _parse_numbered_abcd, _parse_qa_format]:
        result = parser(lines)
        if result:
            return result

    result = _parse_pipe(non_empty)
    if result:
        return result

    return []


def parse_docx(file_path: str) -> list:
    """DOCX fayldan savollarni o'qiydi."""
    try:
        from docx import Document
    except ImportError:
        return []

    try:
        doc = Document(file_path)
    except Exception:
        return []

    questions = []

    def add_q(question, options, answer):
        if question and len(options) >= 2 and answer in options:
            opts = options[:4]
            while len(opts) < 4:
                opts.append(f"Variant {len(opts) + 1}")
            questions.append({"question": question, "options": opts, "answer": answer})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            cells = list(dict.fromkeys(cells))
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)

        if not rows:
            continue

        if all(len(r) == 1 for r in rows):
            flat = [r[0] for r in rows]
            for i in range(0, len(flat), 5):
                b = flat[i : i + 5]
                if len(b) == 5:
                    add_q(b[0], b[1:], b[1])
            continue

        for r in rows:
            if len(r) >= 5:
                add_q(r[0], r[1:5], r[1])
            elif len(r) == 3:
                add_q(r[0], [r[1], r[2], "Variant C", "Variant D"], r[1])

    if not questions:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for parser in [_parse_hash_format, _parse_numbered_abcd]:
            result = parser(lines)
            if result:
                return result
        for i in range(0, len(lines), 5):
            b = lines[i : i + 5]
            if len(b) == 5:
                add_q(b[0], b[1:], b[1])

    return questions


def parse_xlsx(file_path: str) -> list:
    """Excel fayldan savollarni o'qiydi."""
    try:
        import openpyxl
    except ImportError:
        return []

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        questions = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if len(cells) >= 5:
                    q, *opts = cells[:5]
                    questions.append({"question": q, "options": opts[:4], "answer": opts[0]})
        wb.close()
        return questions
    except Exception:
        return []


def parse_pdf(file_path: str) -> list:
    """PDF fayldan savollarni o'qiydi."""
    try:
        import pdfplumber
    except ImportError:
        return []

    try:
        lines = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())

        lines = [l.strip() for l in lines if l.strip()]
        if not lines:
            return []

        for parser in [_parse_hash_format, _parse_numbered_abcd, _parse_pipe, _parse_qa_format]:
            result = parser(lines)
            if result:
                return result
        return []
    except Exception:
        return []


def convert_doc_to_docx(doc_path: str, docx_path: str) -> str:
    """DOC faylni DOCX ga aylantiradi (LibreOffice orqali)."""
    import shutil
    import subprocess

    candidates = [
        "soffice",
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for soffice in candidates:
        found = shutil.which(soffice) or (os.path.isabs(soffice) and os.path.exists(soffice))
        if found:
            out_dir = os.path.dirname(docx_path) or "."
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, doc_path],
                check=True,
                capture_output=True,
                timeout=60,
            )
            auto_out = os.path.join(
                out_dir,
                os.path.splitext(os.path.basename(doc_path))[0] + ".docx",
            )
            if os.path.exists(auto_out) and auto_out != docx_path:
                os.rename(auto_out, docx_path)
            return docx_path

    raise RuntimeError(
        "LibreOffice topilmadi! .doc faylni ochish uchun o'rnating:\n"
        "Linux: sudo apt install libreoffice\n"
        "Windows: https://www.libreoffice.org"
    )


# ─────────────────────────── FSM ─────────────────────────────────────────────

class TestStates(StatesGroup):
    setting_count = State()
    testing = State()


# ─────────────────────────── YORDAMCHI ───────────────────────────────────────

def clear_user_test_session(uid: int):
    key = str(uid)
    if key not in users:
        return
    for k in [
        "selected_questions", "total_test", "current_index", "score",
        "answers", "poll_ids", "waiting_for_skip", "current_poll_message_id",
        "current_poll_id", "current_question_index", "current_answer_recorded",
        "test_start_time",
    ]:
        users[key].pop(k, None)


async def safe_delete(chat_id: int, message_id: int):
    try:
        await bot.delete_message(chat_id, message_id)
    except Exception:
        pass


async def clean_chat(uid: int, chat_id: int, keep_last: int = 0):
    msgs = user_messages.get(uid, [])
    to_del = msgs[:-keep_last] if keep_last else msgs[:]
    for mid in to_del:
        await safe_delete(chat_id, mid)
    user_messages[uid] = msgs[-keep_last:] if keep_last else []


async def add_msg(uid: int, message_id: int):
    user_messages.setdefault(uid, [])
    if message_id not in user_messages[uid]:
        user_messages[uid].append(message_id)


def _norm(text: str, limit: int = 100) -> str:
    t = str(text).strip()
    return (t[: limit - 1] + "…") if len(t) > limit else t


# ─────────────────────────── KLAVIATURALAR ───────────────────────────────────

def main_kb(uid=None):
    row1 = [KeyboardButton(text="📊 Test natijam"), KeyboardButton(text="🆘 Yordam")]
    row2 = [KeyboardButton(text="📄 Yangi test"), KeyboardButton(text="📁 Fayllarim")]
    row3 = [KeyboardButton(text="⭐ Statistika"), KeyboardButton(text="⚙️ Sozlamalar")]
    key = str(uid) if uid else None
    if key and users.get(key, {}).get("uploaded_docs"):
        row2.append(KeyboardButton(text="🔁 Qayta boshlash"))
    return ReplyKeyboardMarkup(keyboard=[row1, row2, row3], resize_keyboard=True)


def count_kb(total: int):
    opts = sorted({n for n in (5, 10, 15, 20, 25, 30, 40, 50) if n <= total} | {total})
    buttons, row = [], []
    for i, c in enumerate(opts):
        label = f"📚 Hammasi ({c})" if c == total else f"📝 {c} ta"
        row.append(InlineKeyboardButton(text=label, callback_data=f"count_{c}"))
        if len(row) == 2 or i == len(opts) - 1:
            buttons.append(row)
            row = []
    buttons.append([InlineKeyboardButton(text="✍️ O'zim kiritaman", callback_data="custom_count")])
    buttons.append([InlineKeyboardButton(text="🎲 Tasodifiy", callback_data="random_count")])
    buttons.append([InlineKeyboardButton(text="❌ Bekor qilish", callback_data="cancel_test")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)


def poll_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⏭ Keyingi savol", callback_data="skip_question")],
        [
            InlineKeyboardButton(text="💡 Javobni ko'rsat", callback_data="show_answer"),
            InlineKeyboardButton(text="🔴 Yakunlash", callback_data="stop_test"),
        ],
    ])


# ─────────────────────────── WEB SERVER ──────────────────────────────────────

async def health(request):
    """Render health check uchun."""
    return web.Response(
        text=json.dumps({"status": "ok", "time": get_uz_time_str()}),
        content_type="application/json",
    )


async def start_web():
    app = web.Application()
    app.router.add_get("/", health)
    app.router.add_get("/health", health)

    runner = web.AppRunner(app)
    await runner.setup()

    port = int(os.environ.get("PORT", 10000))
    site = web.TCPSite(runner, "0.0.0.0", port)
    await site.start()
    print(f"🌐 Web server port {port} da ishga tushdi")


# ─────────────────────────── HANDLER: START ──────────────────────────────────

@dp.message(CommandStart())
async def cmd_start(message: Message, state: FSMContext):
    await state.clear()
    uid = message.from_user.id
    await clean_chat(uid, message.chat.id)
    get_user(uid)  # Yangi foydalanuvchi yaratadi yoki mavjudini qaytaradi

    msg = await message.answer(
        "🎯 <b>TEST MASTER BOT</b>\n"
        "━━━━━━━━━━━━━━━━━━━\n\n"
        f"👋 Assalomu alaykum, <b>{message.from_user.first_name}</b>!\n\n"
        "✨ <b>Imkoniyatlar:</b>\n"
        "• 🎲 Savollar va variantlar aralash\n"
        "• 📝 Test sonini o'zingiz belgilaysiz\n"
        "• 📊 Batafsil statistika\n"
        "• 📁 Fayllar tarixi (qayta ishlating!)\n\n"
        "📎 <b>Qo'llab-quvvatlanadigan formatlar:</b>\n"
        "• <code>.txt</code> — Matnli fayl\n"
        "• <code>.docx</code> — Word\n"
        "• <code>.xlsx</code> — Excel\n"
        "• <code>.pdf</code> — PDF\n\n"
        "💬 Murojaat: @Rustamov_v1",
        parse_mode="HTML",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)


# ─────────────────────────── HANDLER: MENYU ──────────────────────────────────

@dp.message(F.text == "📄 Yangi test")
async def new_test(message: Message, state: FSMContext):
    await state.clear()
    uid = message.from_user.id
    await clean_chat(uid, message.chat.id)
    msg = await message.answer(
        "📄 <b>YANGI TEST</b>\n"
        "━━━━━━━━━━━━━━━━━━━\n\n"
        "📎 Fayl yuboring\n\n"
        "📝 <b>TXT Format 1</b> (<code>#</code> belgisi):\n"
        "<code># Savol matni\n+ To'g'ri javob\n- Noto'g'ri 1\n- Noto'g'ri 2\n- Noto'g'ri 3</code>\n\n"
        "📝 <b>TXT Format 2</b> (A/B/C/D):\n"
        "<code>1. Savol matni\nA) To'g'ri\nB) Noto'g'ri 1\nC) Noto'g'ri 2\nD) Noto'g'ri 3\nJavob: A</code>\n\n"
        "📝 <b>TXT Format 3</b> (Pipe):\n"
        "<code>Savol|A variant|B variant|C variant|D variant</code>\n\n"
        "📘 <b>DOCX:</b> Jadval (1-ustun: Savol, 2–5: Variantlar)\n"
        "📊 <b>XLSX:</b> Jadval (A: Savol, B–E: Variantlar)\n"
        "📄 <b>PDF:</b> Yuqoridagi TXT formatlardan biri",
        parse_mode="HTML",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "🆘 Yordam")
async def cmd_help(message: Message):
    uid = message.from_user.id
    await clean_chat(uid, message.chat.id)
    msg = await message.answer(
        "🆘 <b>YORDAM</b>\n"
        "━━━━━━━━━━━━━━━━━━━\n\n"
        "📌 <b>Foydalanish tartibi:</b>\n"
        "1️⃣ Fayl yuboring (.txt / .docx / .xlsx / .pdf)\n"
        "2️⃣ Test sonini tanlang\n"
        "3️⃣ Javob bering → ⏭ tugmasini bosing\n\n"
        "─────────────────────\n"
        "📝 <b>TXT Format (<code>#</code> belgisi):</b>\n"
        "<code># Savol\n+ To'g'ri\n- Noto'g'ri 1\n- Noto'g'ri 2\n- Noto'g'ri 3</code>\n\n"
        "📝 <b>TXT Format (A/B/C/D):</b>\n"
        "<code>1. Savol\nA) To'g'ri\nB) Noto'g'ri 1\nC) Noto'g'ri 2\nD) Noto'g'ri 3\nJavob: A</code>\n\n"
        "📘 <b>DOCX:</b> 5 ustunli jadval\n"
        "📊 <b>XLSX:</b> 5 ustunli jadval\n\n"
        "⚠️ Har javobdan keyin ⏭ tugmasini bosing!\n\n"
        "💬 Muammo? @Rustamov_v1",
        parse_mode="HTML",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "📊 Test natijam")
async def test_result(message: Message):
    uid = message.from_user.id
    await clean_chat(uid, message.chat.id)
    results = users.get(str(uid), {}).get("results", [])

    if not results:
        msg = await message.answer(
            "📊 <b>TEST NATIJALARI</b>\n"
            "━━━━━━━━━━━━━━━━━━━\n\n"
            "❌ Hali natija yo'q.\n\n📎 Yangi test boshlang!",
            parse_mode="HTML",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)
        return

    text = "📊 <b>TEST NATIJALARI (so'nggi 10)</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
    for i, r in enumerate(results[-10:], 1):
        text += (
            f"<b>{i}.</b> {r['date']}\n"
            f"   📝 {r['total']} ta | ✅ {r['score']} ta\n"
            f"   📈 {r['percentage']:.1f}% | {r['grade']}\n\n"
        )
    msg = await message.answer(text, parse_mode="HTML", reply_markup=main_kb(uid))
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "⭐ Statistika")
async def my_stats(message: Message):
    uid = message.from_user.id
    d = users.get(str(uid), {})
    results = d.get("results", [])
    avg = sum(r["percentage"] for r in results) / len(results) if results else 0
    best = max(results, key=lambda x: x["percentage"]) if results else None

    text = (
        "⭐ <b>UMUMIY STATISTIKA</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        f"📅 Birinchi foydalanish: <b>{d.get('first_visit', 'Noma\'lum')}</b>\n\n"
        f"📊 Jami testlar: <b>{d.get('total_tests', 0)} ta</b>\n"
        f"📝 Jami savollar: <b>{d.get('total_questions', 0)} ta</b>\n"
        f"✅ To'g'ri javoblar: <b>{d.get('total_correct', 0)} ta</b>\n"
        f"📈 O'rtacha natija: <b>{avg:.1f}%</b>\n"
    )
    if best:
        text += f"\n🏆 <b>Eng yaxshi:</b> {best['percentage']:.1f}% — {best['date']}\n"

    msg = await message.answer(text, parse_mode="HTML", reply_markup=main_kb(uid))
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "📁 Fayllarim")
async def my_files(message: Message):
    uid = message.from_user.id
    docs = users.get(str(uid), {}).get("uploaded_docs", [])

    if not docs:
        msg = await message.answer(
            "📁 <b>Fayllarim</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
            "❌ Hali hech qanday fayl yuklanmagan.\n\n📎 Fayl yuboring!",
            parse_mode="HTML",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)
        return

    text = "📁 <b>Mening fayllarim</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
    last_docs = docs[-10:]
    for i, doc in enumerate(last_docs, 1):
        q_count = len(doc.get("questions", []))
        text += f"<b>{i}.</b> {doc['file_name']} — <b>{q_count} savol</b>\n   📅 {doc.get('uploaded_at', '')}\n\n"

    buttons = [
        [InlineKeyboardButton(
            text=f"📄 {d['file_name'][:35]}",
            callback_data=f"selfile_{i}",
        )]
        for i, d in enumerate(docs[-5:])
    ]
    buttons.append([InlineKeyboardButton(text="🏠 Menyu", callback_data="main_menu")])

    msg = await message.answer(
        text,
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
    )
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "⚙️ Sozlamalar")
async def settings_menu(message: Message):
    uid = message.from_user.id
    msg = await message.answer(
        "⚙️ <b>SOZLAMALAR</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        "🔄 Savollar aralashtirish: ✅\n"
        "🔄 Variantlar aralashtirish: ✅\n\n"
        "<i>Tez orada qo'shimcha sozlamalar...</i>",
        parse_mode="HTML",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)


@dp.message(F.text == "🔁 Qayta boshlash")
async def restart_list(message: Message):
    uid = message.from_user.id
    docs = users.get(str(uid), {}).get("uploaded_docs", [])
    if not docs:
        msg = await message.answer(
            "❌ Oldingi fayl topilmadi. Avval fayl yuboring!",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)
        return

    buttons = [
        [InlineKeyboardButton(
            text=f"{i + 1}. {d['file_name']} ({len(d.get('questions', []))} ta)",
            callback_data=f"restart_{i}",
        )]
        for i, d in enumerate(docs[-5:])
    ]
    buttons.append([InlineKeyboardButton(text="❌ Bekor", callback_data="cancel_test")])
    await clean_chat(uid, message.chat.id)
    msg = await message.answer(
        "🔁 <b>OLDINGI FAYLLAR</b>\n━━━━━━━━━━━━━━━━━━━\n\nBirini tanlang:",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
    )
    await add_msg(uid, msg.message_id)


# ─────────────────────────── CALLBACK: FAYL TANLASH ──────────────────────────

@dp.callback_query(F.data.startswith("selfile_"))
async def select_file_cb(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    docs = users.get(str(uid), {}).get("uploaded_docs", [])
    try:
        idx = int(callback.data.split("_")[1])
        doc = docs[idx]
    except (ValueError, IndexError):
        await callback.answer("❌ Xatolik! Fayl topilmadi.", show_alert=True)
        return

    users[str(uid)].update({
        "questions": doc["questions"],
        "total_questions": len(doc["questions"]),
        "file_name": doc["file_name"],
    })
    await state.set_state(TestStates.setting_count)
    await _show_count_selection(callback, uid, len(doc["questions"]))


@dp.callback_query(F.data.startswith("restart_"))
async def restart_doc_cb(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    docs = users.get(str(uid), {}).get("uploaded_docs", [])
    try:
        idx = int(callback.data.split("_")[1])
        doc = docs[idx]
    except (ValueError, IndexError):
        await callback.answer("❌ Noto'g'ri tanlov!", show_alert=True)
        return

    users[str(uid)].update({
        "questions": doc["questions"],
        "total_questions": len(doc["questions"]),
        "file_name": doc["file_name"],
    })
    await state.set_state(TestStates.setting_count)
    await _show_count_selection(callback, uid, len(doc["questions"]))


async def _show_count_selection(callback: CallbackQuery, uid: int, total: int):
    txt = (
        f"📝 <b>TEST SONINI TANLANG</b>\n"
        f"━━━━━━━━━━━━━━━━━━━\n\n"
        f"📚 Mavjud savollar: <b>{total} ta</b>\n\n"
        f"<i>Variantni tanlang 👇</i>"
    )
    try:
        await callback.message.edit_text(txt, parse_mode="HTML", reply_markup=count_kb(total))
    except Exception:
        msg = await callback.message.answer(txt, parse_mode="HTML", reply_markup=count_kb(total))
        await add_msg(uid, msg.message_id)
    await callback.answer()


# ─────────────────────────── HANDLER: FAYL QABUL ─────────────────────────────

@dp.message(F.document)
async def handle_document(message: Message, state: FSMContext):
    uid = message.from_user.id
    await clean_chat(uid, message.chat.id)
    doc = message.document
    fname = (doc.file_name or "fayl").lower()

    ext = next((e for e in SUPPORTED_EXT if fname.endswith(e)), None)
    if ext is None:
        msg = await message.answer(
            "❌ <b>Qo'llab-quvvatlanmaydigan fayl!</b>\n\n"
            "✅ Qabul qilinadi:\n"
            "• <code>.txt</code>  — Matnli fayl\n"
            "• <code>.docx</code> — Word\n"
            "• <code>.doc</code>  — Eski Word\n"
            "• <code>.xlsx</code> — Excel\n"
            "• <code>.pdf</code>  — PDF",
            parse_mode="HTML",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)
        return

    loading = await message.answer(
        "⏳ <b>Fayl yuklanmoqda...</b>\n━━━━━━━━━━━━━━━━━━━\n🔄 Tahlil qilinmoqda...",
        parse_mode="HTML",
    )
    await add_msg(uid, loading.message_id)

    save_path = None
    try:
        cleanup_old_files()
        tg_file = await bot.get_file(doc.file_id)
        downloaded = await bot.download_file(tg_file.file_path)
        save_path = os.path.join("temp", f"u{uid}_{int(time.time())}{ext}")
        with open(save_path, "wb") as f:
            f.write(downloaded.read())

        # .doc → .docx konversiya
        parse_ext = ext
        if ext == ".doc":
            converted = save_path.replace(".doc", "_conv.docx")
            save_path = convert_doc_to_docx(save_path, converted)
            parse_ext = ".docx"

        # Parse qilish
        fmt_map = {
            ".txt": (parse_txt, "TXT"),
            ".docx": (parse_docx, "DOCX"),
            ".xlsx": (parse_xlsx, "XLSX"),
            ".pdf": (parse_pdf, "PDF"),
        }
        parser_fn, fmt_name = fmt_map.get(parse_ext, (lambda p: [], "?"))
        questions = parser_fn(save_path)

        await clean_chat(uid, message.chat.id)

        if not questions:
            hint = (
                "📌 <b>TXT uchun to'g'ri format:</b>\n"
                "<code># Savol matni\n+ To'g'ri javob\n- Noto'g'ri 1\n- Noto'g'ri 2\n- Noto'g'ri 3</code>\n\n"
                "yoki\n\n"
                "<code>1. Savol\nA) To'g'ri\nB) Noto'g'ri 1\nC) Noto'g'ri 2\nD) Noto'g'ri 3\nJavob: A</code>"
                if parse_ext == ".txt"
                else "📘 <b>DOCX/XLSX uchun:</b> 5 ustunli jadval kerak\n(1-ustun: Savol, 2-5-ustunlar: Variantlar)"
            )
            msg = await message.answer(
                f"❌ <b>Savol topilmadi!</b>\n\n"
                f"<i>{fmt_name} fayldan savol o'qib bo'lmadi.</i>\n\n{hint}",
                parse_mode="HTML",
                reply_markup=main_kb(uid),
            )
            await add_msg(uid, msg.message_id)
            return

        # Saqlash
        user_data = get_user(uid)
        uploaded_docs = user_data.get("uploaded_docs", [])
        uploaded_docs.append({
            "file_name": doc.file_name,
            "file_path": save_path,
            "questions": questions,
            "uploaded_at": get_uz_time_str(),
        })
        # Maksimal 20 ta fayl saqlash
        if len(uploaded_docs) > 20:
            uploaded_docs = uploaded_docs[-20:]

        user_data.update({
            "questions": questions,
            "total_questions": len(questions),
            "file_name": doc.file_name,
            "uploaded_file": save_path,
            "uploaded_docs": uploaded_docs,
        })
        save_user(uid)

        await state.set_state(TestStates.setting_count)

        msg = await message.answer(
            f"✅ <b>Fayl muvaffaqiyatli yuklandi!</b>\n"
            f"━━━━━━━━━━━━━━━━━━━\n\n"
            f"📋 Format: <b>{fmt_name}</b>\n"
            f"📚 Jami savollar: <b>{len(questions)} ta</b>\n"
            f"📄 Fayl: <code>{doc.file_name}</code>\n\n"
            f"<i>Test sonini tanlang 👇</i>",
            parse_mode="HTML",
            reply_markup=count_kb(len(questions)),
        )
        await add_msg(uid, msg.message_id)

    except Exception as e:
        await clean_chat(uid, message.chat.id)
        err_text = str(e)[:300]
        msg = await message.answer(
            f"❌ <b>Xatolik yuz berdi:</b>\n<code>{err_text}</code>\n\n"
            "Iltimos, fayl formatini tekshiring va qaytadan urinib ko'ring.",
            parse_mode="HTML",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)


# ─────────────────────────── HANDLER: TEST SOZLASH ───────────────────────────

@dp.callback_query(F.data.startswith("count_"))
async def select_count(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    key = str(uid)
    if key not in users or "total_questions" not in users[key]:
        await callback.answer("❌ Avval fayl yuklang!", show_alert=True)
        return

    try:
        count = int(callback.data.split("_")[1])
    except (ValueError, IndexError):
        await callback.answer("❌ Noto'g'ri qiymat!", show_alert=True)
        return

    await safe_delete(callback.message.chat.id, callback.message.message_id)
    await start_test(callback.message, uid, count, state)
    await callback.answer(f"✅ {count} ta test boshlandi!")


@dp.callback_query(F.data == "random_count")
async def random_count_cb(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    key = str(uid)
    if key not in users or "total_questions" not in users[key]:
        await callback.answer("❌ Avval fayl yuklang!", show_alert=True)
        return

    total = users[key]["total_questions"]
    count = random.randint(min(5, total), min(50, total))
    await safe_delete(callback.message.chat.id, callback.message.message_id)
    await start_test(callback.message, uid, count, state)
    await callback.answer(f"🎲 {count} ta test boshlandi!")


@dp.callback_query(F.data == "custom_count")
async def custom_count_prompt(callback: CallbackQuery, state: FSMContext):
    await state.set_state(TestStates.setting_count)
    try:
        await callback.message.edit_text(
            "✍️ <b>TEST SONINI KIRITING</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
            "Nechta savol bo'lishini raqam bilan yozing:\n"
            "<i>Masalan: 15, 25, 30...</i>",
            parse_mode="HTML",
        )
    except Exception:
        pass
    await callback.answer()


@dp.message(TestStates.setting_count)
async def process_custom_count(message: Message, state: FSMContext):
    uid = message.from_user.id
    key = str(uid)

    if message.text and message.text.startswith("/"):
        return

    await clean_chat(uid, message.chat.id)

    if key not in users or "total_questions" not in users[key]:
        msg = await message.answer("❌ Avval fayl yuklang!", reply_markup=main_kb(uid))
        await add_msg(uid, msg.message_id)
        await state.clear()
        return

    try:
        count = int(message.text.strip())
        total = users[key]["total_questions"]
        if count < 1 or count > total:
            msg = await message.answer(
                f"❌ 1 dan {total} gacha raqam kiriting",
                reply_markup=main_kb(uid),
            )
            await add_msg(uid, msg.message_id)
            return
        await start_test(message, uid, count, state)
    except (ValueError, TypeError):
        msg = await message.answer(
            "❌ Faqat raqam kiriting! (Masalan: 10)",
            reply_markup=main_kb(uid),
        )
        await add_msg(uid, msg.message_id)


# ─────────────────────────── TEST JARAYONI ───────────────────────────────────

async def start_test(message: Message, uid: int, count: int, state: FSMContext):
    await clean_chat(uid, message.chat.id)
    key = str(uid)

    pool = copy.deepcopy(users[key]["questions"])
    random.shuffle(pool)
    selected = pool[:count]
    for q in selected:
        random.shuffle(q["options"])

    users[key].update({
        "selected_questions": selected,
        "total_test": count,
        "current_index": 0,
        "score": 0,
        "answers": [],
        "poll_ids": [],
        "waiting_for_skip": False,
        "current_answer_recorded": False,
        "current_poll_message_id": None,
        "current_poll_id": None,
        "current_question_index": 0,
        "test_start_time": get_uz_time_str(),
    })
    await state.set_state(TestStates.testing)

    msg = await message.answer(
        f"🚀 <b>TEST BOSHLANDI!</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        f"📝 Jami: <b>{count} ta savol</b>\n"
        f"🎲 Savollar aralash holda\n\n"
        f"⚠️ <b>Muhim:</b> Javobdan keyin ⏭ tugmasini bosing!\n\n"
        f"<i>Omad! 🍀</i>",
        parse_mode="HTML",
    )
    await add_msg(uid, msg.message_id)
    await asyncio.sleep(1)
    await clean_chat(uid, message.chat.id)
    await send_poll(message.chat.id, uid)


async def send_poll(chat_id: int, uid: int):
    key = str(uid)
    data = users.get(key)
    if not data:
        return

    idx = data.get("current_index", 0)
    selected = data.get("selected_questions", [])

    if idx >= len(selected):
        return

    qd = selected[idx]
    opts = [_norm(o) for o in qd["options"]]
    ans = _norm(qd["answer"])

    # To'g'ri javob indeksini topish
    try:
        correct_id = opts.index(ans)
    except ValueError:
        # Agar topilmasa, birinchi variantni to'g'ri deb belgilash
        correct_id = 0
        if opts:
            opts[0] = ans
        else:
            opts = [ans, "Variant B", "Variant C", "Variant D"]

    q_text = str(qd["question"]).strip()
    if len(q_text) > 300:
        q_text = q_text[:297] + "..."

    # Variantlar uzunligini tekshirish (Telegram max 100 belgi)
    opts = [o[:99] for o in opts]

    # Oldingi poll xabarini o'chirish
    prev_msg_id = data.get("current_poll_message_id")
    if prev_msg_id:
        await safe_delete(chat_id, prev_msg_id)
    await clean_chat(uid, chat_id)

    try:
        poll_msg = await bot.send_poll(
            chat_id=chat_id,
            question=f"📝 {idx + 1}/{data['total_test']}\n\n{q_text}",
            options=opts,
            type="quiz",
            correct_option_id=correct_id,
            explanation=f"✅ To'g'ri: {ans[:200]}",
            is_anonymous=False,
            reply_markup=poll_kb(),
        )
        data.update({
            "poll_ids": data.get("poll_ids", []) + [poll_msg.poll.id],
            "current_poll_id": poll_msg.poll.id,
            "current_question_index": idx,
            "current_poll_message_id": poll_msg.message_id,
            "waiting_for_skip": True,
            "current_answer_recorded": False,
        })
    except Exception as e:
        print(f"❌ Poll yuborishda xatolik: {e}")
        try:
            err_msg = await bot.send_message(
                chat_id,
                f"❌ Savol yuborishda xatolik yuz berdi.\n⏭ Keyingisiga o'tish uchun tugmani bosing.",
                reply_markup=poll_kb(),
            )
            data["current_poll_message_id"] = err_msg.message_id
            data["waiting_for_skip"] = True
        except Exception:
            pass


@dp.poll_answer()
async def on_poll_answer(poll_answer):
    uid = poll_answer.user.id
    key = str(uid)
    data = users.get(key)

    if not data or "selected_questions" not in data:
        return
    if poll_answer.poll_id != data.get("current_poll_id"):
        return
    if data.get("current_answer_recorded"):
        return
    if not poll_answer.option_ids:
        return

    cidx = data.get("current_question_index", 0)
    selected = data.get("selected_questions", [])

    if cidx >= len(selected):
        return

    qd = selected[cidx]
    try:
        chosen = qd["options"][poll_answer.option_ids[0]]
    except IndexError:
        return

    correct = chosen == qd["answer"]
    data.setdefault("answers", []).append({
        "question": qd["question"],
        "user_answer": chosen,
        "correct_answer": qd["answer"],
        "is_correct": correct,
    })
    if correct:
        data["score"] = data.get("score", 0) + 1
    data["current_answer_recorded"] = True


@dp.callback_query(F.data == "show_answer")
async def show_answer(callback: CallbackQuery):
    uid = callback.from_user.id
    key = str(uid)
    data = users.get(key)

    if not data or "selected_questions" not in data:
        await callback.answer("❌ Test topilmadi!", show_alert=True)
        return

    cidx = data.get("current_question_index", 0)
    selected = data.get("selected_questions", [])

    if cidx >= len(selected):
        await callback.answer("❌ Savol topilmadi!", show_alert=True)
        return

    qd = selected[cidx]
    answer_text = str(qd["answer"])[:200]
    await callback.answer(f"✅ To'g'ri javob:\n{answer_text}", show_alert=True)


@dp.callback_query(F.data == "skip_question")
async def skip_question(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    key = str(uid)
    data = users.get(key)

    if not data:
        await callback.answer("❌ Test topilmadi!", show_alert=True)
        return
    if not data.get("waiting_for_skip"):
        await callback.answer("⏳ Avval javob bering!", show_alert=True)
        return

    cidx = data.get("current_question_index", 0)
    selected = data.get("selected_questions", [])

    # Javob berilmagan bo'lsa, noto'g'ri deb belgilash
    if not data.get("current_answer_recorded") and cidx < len(selected):
        qd = selected[cidx]
        data.setdefault("answers", []).append({
            "question": qd["question"],
            "user_answer": "Javob berilmadi",
            "correct_answer": qd["answer"],
            "is_correct": False,
        })

    data["waiting_for_skip"] = False
    data["current_answer_recorded"] = False
    data["current_index"] = data.get("current_index", 0) + 1

    total_test = data.get("total_test", 0)
    if data["current_index"] >= total_test:
        await clean_chat(uid, callback.message.chat.id)
        await safe_delete(callback.message.chat.id, callback.message.message_id)
        await show_results(callback.message.chat.id, uid)
        await state.clear()
        await callback.answer("✅ Test yakunlandi!")
        return

    await callback.answer("⏭ Keyingi...")
    await clean_chat(uid, callback.message.chat.id)
    await send_poll(callback.message.chat.id, uid)


@dp.callback_query(F.data == "stop_test")
async def stop_test(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    key = str(uid)
    data = users.get(key)

    if not data:
        await callback.answer("❌ Test topilmadi!", show_alert=True)
        return

    cidx = data.get("current_question_index", 0)
    answers = data.get("answers", [])
    selected = data.get("selected_questions", [])

    # Joriy savolga javob berilmagan bo'lsa
    if len(answers) <= cidx < len(selected):
        qd = selected[cidx]
        answers.append({
            "question": qd["question"],
            "user_answer": "Test yakunlandi",
            "correct_answer": qd["answer"],
            "is_correct": False,
        })
        data["answers"] = answers

    await clean_chat(uid, callback.message.chat.id)
    await safe_delete(callback.message.chat.id, callback.message.message_id)
    await show_results(callback.message.chat.id, uid, stopped=True)
    await state.clear()
    await callback.answer("🔴 Test yakunlandi")


# ─────────────────────────── NATIJALAR ───────────────────────────────────────

async def show_results(chat_id: int, uid: int, stopped: bool = False):
    key = str(uid)
    data = users.get(key, {})
    score = data.get("score", 0)
    total = data.get("total_test", 0)
    answered = len(data.get("answers", []))
    pct = (score / total * 100) if total else 0

    if pct >= 90:
        grade, emoji = "A'lo", "🏆"
    elif pct >= 75:
        grade, emoji = "Yaxshi", "🎉"
    elif pct >= 60:
        grade, emoji = "Qoniqarli", "👍"
    else:
        grade, emoji = "O'qish kerak", "📚"

    # Vaqt hisoblash
    start_str = data.get("test_start_time")
    time_str = ""
    if start_str:
        try:
            start = datetime.strptime(start_str, "%d.%m.%Y %H:%M").replace(tzinfo=UZBEKISTAN_TZ)
            diff = get_uz_time() - start
            total_secs = int(diff.total_seconds())
            m, s = total_secs // 60, total_secs % 60
            time_str = f"\n⏱ Vaqt: {m} min {s} sek"
        except Exception:
            pass

    text = (
        f"{emoji} <b>TEST NATIJASI</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        f"📊 <b>Statistika:</b>\n"
        f"• Jami: <b>{total} ta</b>\n"
        f"• Javob berilgan: <b>{answered} ta</b>\n"
        f"• O'tkazilgan: <b>{max(total - answered, 0)} ta</b>{time_str}\n\n"
        f"✅ To'g'ri: <b>{score} ta</b>\n"
        f"❌ Noto'g'ri: <b>{max(total - score, 0)} ta</b>\n"
        f"📈 Foiz: <b>{pct:.1f}%</b>\n"
        f"🏆 Baho: <b>{grade}</b>\n\n"
        f"<i>{'⚠️ Test vaqtidan oldin yakunlandi' if stopped else '🎊 Test muvaffaqiyatli yakunlandi!'}</i>"
    )

    # Statistikani yangilash
    data["total_tests"] = data.get("total_tests", 0) + 1
    data["total_questions"] = data.get("total_questions", 0) + total
    data["total_correct"] = data.get("total_correct", 0) + score
    data.setdefault("results", []).append({
        "date": get_uz_time_str(),
        "total": total,
        "score": score,
        "percentage": pct,
        "grade": grade,
    })
    # Maksimal 100 ta natija saqlash
    if len(data["results"]) > 100:
        data["results"] = data["results"][-100:]

    save_user(uid)

    await clean_chat(uid, chat_id)
    msg = await bot.send_message(
        chat_id,
        text,
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="🔄 Qayta urinish", callback_data="retry_poll"),
                InlineKeyboardButton(text="📋 Batafsil", callback_data="poll_details"),
            ],
            [
                InlineKeyboardButton(text="📊 Barcha natijalar", callback_data="all_results"),
                InlineKeyboardButton(text="🏠 Menyu", callback_data="main_menu"),
            ],
        ]),
    )
    await add_msg(uid, msg.message_id)


@dp.callback_query(F.data == "retry_poll")
async def retry_poll(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    key = str(uid)
    data = users.get(key)

    if not data or "selected_questions" not in data:
        await callback.answer("❌ Ma'lumot topilmadi. Yangi fayl yuklang.", show_alert=True)
        return

    selected = copy.deepcopy(data["selected_questions"])
    random.shuffle(selected)
    for q in selected:
        random.shuffle(q["options"])

    data.update({
        "selected_questions": selected,
        "current_index": 0,
        "score": 0,
        "answers": [],
        "poll_ids": [],
        "waiting_for_skip": False,
        "current_answer_recorded": False,
        "current_poll_message_id": None,
        "current_poll_id": None,
        "current_question_index": 0,
        "test_start_time": get_uz_time_str(),
    })
    await state.set_state(TestStates.testing)
    await clean_chat(uid, callback.message.chat.id)
    await safe_delete(callback.message.chat.id, callback.message.message_id)

    msg = await callback.message.answer(
        f"🔄 <b>TEST QAYTA BOSHLANDI!</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        f"📝 {data['total_test']} ta savol | 🎲 Aralashtirildi\n\n<i>Omad! 🍀</i>",
        parse_mode="HTML",
    )
    await add_msg(uid, msg.message_id)
    await asyncio.sleep(1)
    await clean_chat(uid, callback.message.chat.id)
    await send_poll(callback.message.chat.id, uid)
    await callback.answer("✅ Boshlandi")


@dp.callback_query(F.data == "poll_details")
async def poll_details(callback: CallbackQuery):
    uid = callback.from_user.id
    answers = users.get(str(uid), {}).get("answers", [])

    if not answers:
        await callback.answer("❌ Javoblar yo'q", show_alert=True)
        return

    text = "📋 <b>BATAFSIL NATIJALAR</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
    for i, a in enumerate(answers, 1):
        icon = "✅" if a["is_correct"] else "❌"
        q = a["question"][:60] + ("..." if len(a["question"]) > 60 else "")
        text += f"<b>{i}.</b> {q}\n   {icon} {a['user_answer']}\n"
        if not a["is_correct"]:
            text += f"   ✅ To'g'ri: {a['correct_answer']}\n"
        text += "\n"

        # Xabar 3500 belgidan oshsa, bo'lib yuborish
        if len(text) > 3500:
            msg = await callback.message.answer(text, parse_mode="HTML")
            await add_msg(uid, msg.message_id)
            text = ""

    if text:
        msg = await callback.message.answer(text, parse_mode="HTML")
        await add_msg(uid, msg.message_id)

    await callback.answer()


@dp.callback_query(F.data == "all_results")
async def all_results_cb(callback: CallbackQuery):
    uid = callback.from_user.id
    results = users.get(str(uid), {}).get("results", [])

    if not results:
        await callback.answer("❌ Natijalar yo'q", show_alert=True)
        return

    await clean_chat(uid, callback.message.chat.id)
    await safe_delete(callback.message.chat.id, callback.message.message_id)

    text = "📊 <b>BARCHA NATIJALAR</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
    for i, r in enumerate(results[-10:], 1):
        text += (
            f"<b>{i}.</b> {r['date']}\n"
            f"   📝 {r['total']} ta | ✅ {r['score']} ta\n"
            f"   📈 {r['percentage']:.1f}% | {r['grade']}\n\n"
        )
    msg = await callback.message.answer(text, parse_mode="HTML", reply_markup=main_kb(uid))
    await add_msg(uid, msg.message_id)
    await callback.answer()


@dp.callback_query(F.data == "main_menu")
async def back_to_menu(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    uid = callback.from_user.id
    clear_user_test_session(uid)
    await safe_delete(callback.message.chat.id, callback.message.message_id)
    await clean_chat(uid, callback.message.chat.id)
    msg = await bot.send_message(
        callback.message.chat.id,
        "🏠 <b>BOSH MENYU</b>\n━━━━━━━━━━━━━━━━━━━\n\n"
        "📎 Yangi test uchun fayl yuboring 👇",
        parse_mode="HTML",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)
    await callback.answer("🏠 Bosh menyu")


@dp.callback_query(F.data == "cancel_test")
async def cancel_test(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    uid = callback.from_user.id
    clear_user_test_session(uid)
    await clean_chat(uid, callback.message.chat.id)
    await safe_delete(callback.message.chat.id, callback.message.message_id)
    msg = await bot.send_message(
        callback.message.chat.id,
        "❌ Bekor qilindi\n\n📎 Yangi test uchun fayl yuboring",
        reply_markup=main_kb(uid),
    )
    await add_msg(uid, msg.message_id)
    await callback.answer()


# ─────────────────────────── MAIN ────────────────────────────────────────────

async def main():
    print("🚀 Bot ishga tushdi...")
    print(f"📍 Vaqt zonasi: UTC+5 (O'zbekiston)")
    cleanup_old_files()
    await start_web()
    print("✅ Polling boshlandi...")
    await dp.start_polling(bot, allowed_updates=["message", "callback_query", "poll_answer"])


if __name__ == "__main__":
    asyncio.run(main())
